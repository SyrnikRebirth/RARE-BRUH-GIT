import datetime
import time
import sys
import threading
from docx import Document

def turn_txt_to_docx(file_name):
    document = Document()
    document.add_heading(str(datetime.date.today()))
    string_num = 0
    with open(file_name, "r") as file:
        for string in file:
            if (string_num % 2 == 0):
                p = document.add_paragraph()
                p.add_run(string).bold = True
            else:
                p = document.add_paragraph()
                p.add_run(string[:6]).bold = True
                p.add_run(string[6:]).italic = True
            string_num += 1
    document.save(file_name[:-4] + '.docx')
