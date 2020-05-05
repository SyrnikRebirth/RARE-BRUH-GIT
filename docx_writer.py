import datetime
import time
import sys
import threading
from docx import Document

def my_input(stop_event, command):
    while not stop_event.is_set():
        command[0] = input()

def main():
    document = Document()
    document.add_heading(str(datetime.date.today()))
    command = [""]
    s_stop = threading.Event()
    s = threading.Thread(target=my_input, args=(s_stop, command), daemon=True)
    s.start()
    while True:
        while command[0] == "":
            time.sleep(1)
        if command[0] == "stop_file_writing_qwertasdf":
            document.save('test.docx')
            s_stop.set()
            break
        else:
            string = command[0]
            separator = string.find(':')
            p = document.add_paragraph()
            p.add_run(string[:separator]).bold = True
            p.add_run(string[separator:]).italic = True
        command[0] = ""
    s.join()


if __name__ == "__main__":
    main()
